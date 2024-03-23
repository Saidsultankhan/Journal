from django.http import HttpResponse
from rest_framework.views import APIView
from src.apps.jounal.models import DairyOfClass
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert



class WordToPDFView(APIView):
    def get(self, request):
        database = []
        pupils_data = DairyOfClass.objects.filter(pupil__user=request.user).prefetch_related('pupil', 'subject', 'grade')
        user = request.user.pupil_user.values().first()['name_uz']

        for pupil in pupils_data:
            database.append(
                [
                    f"{user}", 
                    f"{str(pupil.grade)}",
                    f"{str(pupil.subject)}",
                    f"{pupil.quarter}",
                    f"{pupil.mark}",
                ]
            )
        data = {
            'headers': ['Name', 'Grade', 'Subject', 'Quarter', 'Mark'],
            'rows': database
        }

        document = Document()
        document.add_paragraph(f'Butun yil davomida olingan {user}-ning baholari.', style='Heading1')

        table = document.add_table(rows=1, cols=len(data['headers']))
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(data['headers']):
            hdr_cells[idx].text = header


        for row in data['rows']:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = value

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        with open('temporary.docx', 'wb') as f:
            f.write(buffer.getvalue())

        output_file = f'{user}-ning baho tabeli.pdf'
        
        with open(output_file, 'wb') as pdf:
            pass

        convert('temporary.docx', output_file)
        
        with open(output_file, 'rb') as pdf_file:
            response = HttpResponse(pdf_file.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename={output_file}'
            return response
